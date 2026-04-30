from flask import Blueprint, render_template, request
from models import db, Student, Book, Circulation
from datetime import datetime

dashboard_bp = Blueprint('dashboard', __name__)

@dashboard_bp.route('/')
def index():
    student_count = Student.query.count()
    book_count = Book.query.count()
    active_issues = Circulation.query.filter(Circulation.return_date == None).count()
    overdue_issues = Circulation.query.filter(
        Circulation.return_date == None, 
        Circulation.due_date < datetime.utcnow()
    ).count()

    recent_issues = Circulation.query.filter(Circulation.return_date == None).order_by(Circulation.issue_date.desc()).limit(5).all()

    return render_template('dashboard/index.html', 
                           student_count=student_count, 
                           book_count=book_count, 
                           active_issues=active_issues, 
                           overdue_issues=overdue_issues,
                           recent_issues=recent_issues,
                           current_time=datetime.utcnow())

@dashboard_bp.route('/reports', methods=['GET', 'POST'])
def reports():
    from flask import make_response
    from xhtml2pdf import pisa
    import io

    if request.method == 'POST':
        report_type = request.form.get('report_type')
        format_type = request.form.get('format_type')
        
        data = []
        headers = []
        title = ""

        if report_type == 'students':
            title = "Comprehensive Student Details"
            students = Student.query.order_by(Student.class_name, Student.division, Student.roll_no).all()
            for s in students:
                data.append({'ID': s.id, 'Roll No': s.roll_no, 'Name': s.name, 'Class': s.class_name, 'Division': s.division})
            headers = ['ID', 'Roll No', 'Name', 'Class', 'Division']

        elif report_type == 'books':
            title = "Comprehensive Book Inventory"
            books = Book.query.order_by(Book.title).all()
            for b in books:
                data.append({'Title': b.title, 'Author': b.author, 'Category': b.category, 'Copies': b.copies, 'Price': b.price})
            headers = ['Title', 'Author', 'Category', 'Copies', 'Price']

        elif report_type == 'returns':
            title = "Library Book Circulation History"
            # Get ALL circulations (both taken and returned)
            circulations = Circulation.query.order_by(Circulation.issue_date.desc()).all()
            for c in circulations:
                data.append({
                    'Student Name': c.student.name,
                    'Class & Div': f"{c.student.class_name} ({c.student.division})",
                    'Roll No': c.student.roll_no,
                    'Book Title': c.book.title,
                    'Issue Date': c.issue_date.strftime('%Y-%m-%d'),
                    'Return Date': c.return_date.strftime('%Y-%m-%d') if c.return_date else 'Not Returned',
                    'Late Days': c.late_days
                })
            headers = ['Student Name', 'Class & Div', 'Roll No', 'Book Title', 'Issue Date', 'Return Date', 'Late Days']

        if not data:
            flash("No data available for the selected report.", "warning")
            return redirect(url_for('dashboard.reports'))

        if format_type == 'csv':
            import csv
            si = io.StringIO()
            writer = csv.DictWriter(si, fieldnames=headers)
            writer.writeheader()
            for row in data:
                writer.writerow(row)
            
            response = make_response(si.getvalue())
            response.headers["Content-Disposition"] = f"attachment; filename=report_{report_type}.csv"
            response.headers["Content-type"] = "text/csv"
            return response

        elif format_type == 'pdf':
            html = render_template('dashboard/pdf_template.html', title=title, headers=headers, data=data)
            result = io.BytesIO()
            pdf = pisa.pisaDocument(io.BytesIO(html.encode("ISO-8859-1")), result)
            if not pdf.err:
                response = make_response(result.getvalue())
                response.headers['Content-Type'] = 'application/pdf'
                response.headers['Content-Disposition'] = f'inline; filename=report_{report_type}.pdf'
                return response
            else:
                flash("Error generating PDF.", "danger")
                
        elif format_type == 'docx':
            from docx import Document
            doc = Document()
            doc.add_heading(title, 0)
            
            # Create a table
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                
            for row_data in data:
                row_cells = table.add_row().cells
                for i, header in enumerate(headers):
                    row_cells[i].text = str(row_data.get(header, ''))
                    
            result = io.BytesIO()
            doc.save(result)
            response = make_response(result.getvalue())
            response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            response.headers['Content-Disposition'] = f'attachment; filename=report_{report_type}.docx'
            return response

    return render_template('dashboard/reports.html')

@dashboard_bp.route('/analytics', methods=['GET', 'POST'])
def analytics():
    from collections import defaultdict
    from flask import make_response
    import csv
    import io

    circulations = db.session.query(Circulation, Student, Book).join(Student).join(Book).all()
    
    monthly_stats = defaultdict(int)
    yearly_stats = defaultdict(int)
    class_stats = defaultdict(int)
    subject_stats = defaultdict(int)
    student_stats = defaultdict(int)
    
    for circ, student, book in circulations:
        month_key = circ.issue_date.strftime('%Y-%m')
        year_key = circ.issue_date.strftime('%Y')
        monthly_stats[month_key] += 1
        yearly_stats[year_key] += 1
        
        class_stats[student.class_name.strip().upper()] += 1
        
        if book.subject:
            normalized_subject = book.subject.strip().title()
            subject_stats[normalized_subject] += 1
            
        student_label = f"{student.name} (Class {student.class_name}, Div {student.division}, Roll {student.roll_no})"
        student_stats[student_label] += 1

    monthly_sorted = dict(sorted(monthly_stats.items(), reverse=False)) # Chronological for chart
    yearly_sorted = dict(sorted(yearly_stats.items(), reverse=False))

    class_sorted = dict(sorted(class_stats.items(), key=lambda item: item[1], reverse=True))
    subject_sorted = dict(sorted(subject_stats.items(), key=lambda item: item[1], reverse=True))
    student_sorted = dict(sorted(student_stats.items(), key=lambda item: item[1], reverse=True)[:50])
    
    if request.method == 'POST':
        format_type = request.form.get('format_type')
        
        if format_type == 'csv':
            si = io.StringIO()
            writer = csv.writer(si)
            writer.writerow(['Issue Date', 'Month', 'Year', 'Student Name', 'Roll No', 'Class', 'Book Title', 'Subject'])
            for circ, student, book in circulations:
                writer.writerow([circ.issue_date.strftime('%Y-%m-%d'), circ.issue_date.strftime('%Y-%m'), circ.issue_date.strftime('%Y'), student.name, student.roll_no, student.class_name, book.title, book.subject])
            response = make_response(si.getvalue())
            response.headers["Content-Disposition"] = "attachment; filename=library_analytics.csv"
            response.headers["Content-type"] = "text/csv"
            return response
            
        elif format_type == 'docx':
            from docx import Document
            doc = Document()
            doc.add_heading('Library Analytics Report', 0)
            
            doc.add_heading('Most Active Classes', level=1)
            for k, v in list(class_sorted.items())[:10]:
                doc.add_paragraph(f"Class {k}: {v} books issued")
                
            doc.add_heading('Most Popular Subjects', level=1)
            for k, v in list(subject_sorted.items())[:10]:
                doc.add_paragraph(f"{k}: {v} books issued")
                
            doc.add_heading('Top Students', level=1)
            for k, v in list(student_sorted.items())[:10]:
                doc.add_paragraph(f"{k}: {v} books issued")
                
            doc.add_heading('Monthly Activity', level=1)
            for k, v in monthly_sorted.items():
                doc.add_paragraph(f"{k}: {v} books issued")
                
            result = io.BytesIO()
            doc.save(result)
            response = make_response(result.getvalue())
            response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            response.headers['Content-Disposition'] = 'attachment; filename=library_analytics.docx'
            return response
            
    return render_template('dashboard/analytics.html', 
                          monthly=monthly_sorted, 
                          yearly=yearly_sorted, 
                          classes=class_sorted, 
                          subjects=subject_sorted, 
                          students=student_sorted)

@dashboard_bp.route('/about')
def about():
    return render_template('dashboard/about.html')
